"""Minimal, evidence-linked DOCX renderer for the step 11 delivery gate."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.app.db.models import Task


ACCENT = RGBColor(86, 125, 31)
INK = RGBColor(43, 49, 53)
CJK_FONT = "STHeiti"


def _text(value: Any, fallback: str = "待确认") -> str:
    if value is None or value == "":
        return fallback
    if isinstance(value, list):
        return "；".join(str(item) for item in value) or fallback
    return str(value)


def _apply_cjk_font(run) -> None:
    run.font.name = CJK_FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), CJK_FONT)


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(cell, value: Any, *, bold: bool = False, color: RGBColor = INK, fill: str | None = None) -> None:
    cell.text = ""
    if fill:
        _shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(_text(value))
    _apply_cjk_font(run)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        _apply_cjk_font(run)
        run.font.color.rgb = ACCENT if level == 1 else INK
    if level == 1:
        properties = paragraph._p.get_or_add_pPr()
        borders = properties.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            properties.append(borders)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "86BC25")
        borders.append(left)


def build_review_docx(task: Task, objects: dict[str, Any], output_path: str | Path) -> dict[str, Any]:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    header = section.header.paragraphs[0]
    header.text = "外规解读智能体工作台  |  仅基于已登记监管原文"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _apply_cjk_font(header.runs[0])
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 110, 116)
    footer = section.footer.paragraphs[0]
    footer.text = "人工复核后版本 · 证据链和待确认边界随报告保留"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_cjk_font(footer.runs[0])
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 110, 116)

    normal = document.styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    for style_name in ("Normal", "Heading 1", "Heading 2", "List Bullet"):
        style = document.styles[style_name]
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), CJK_FONT)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("外规解读智能体工作台")
    _apply_cjk_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = ACCENT
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.left_indent = Inches(0.08)
    subtitle = document.add_paragraph(task.task_name)
    subtitle.paragraph_format.space_after = Pt(2)
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.bold = True
    _apply_cjk_font(subtitle.runs[0])
    meta = document.add_paragraph(f"生成时间：{datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M %Z')}  |  复核状态：已通过质量闸门")
    meta.paragraph_format.space_after = Pt(12)
    meta.runs[0].font.size = Pt(8.5)
    meta.runs[0].font.color.rgb = RGBColor(100, 110, 116)
    _apply_cjk_font(meta.runs[0])

    regulation = task.regulation
    version = next((item for item in regulation.versions if item.is_current), None) if regulation else None
    _add_heading(document, "一、法规概览")
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    metadata = [
        ("法规名称", regulation.title if regulation else None),
        ("文号", regulation.document_no if regulation else None),
        ("发布机关", regulation.issuer if regulation else None),
        ("版本", version.version_label if version else None),
        ("发布日期", version.publish_date if version else None),
        ("生效日期", version.effective_date if version else None),
        ("版本比较", "待补充权威原文；当前不生成差异结论"),
    ]
    for label, value in metadata:
        cells = table.add_row().cells
        _set_cell_text(cells[0], label, bold=True, color=ACCENT, fill="F7F7F8")
        _set_cell_text(cells[1], value)

    s2 = ((task.step_status or {}).get("S2") or {}).get("output") or {}
    _add_heading(document, "二、机构适用性判断")
    paragraph = document.add_paragraph()
    lead = paragraph.add_run(f"结论：{_text(s2.get('status'))}。 ")
    lead.bold = True
    _apply_cjk_font(lead)
    reason = paragraph.add_run(_text(s2.get("reason")))
    _apply_cjk_font(reason)
    paragraph.paragraph_format.space_after = Pt(8)

    _add_heading(document, "三、整体外规解读")
    overall = objects["overall"]
    document.add_paragraph(_text(overall.summary, "待确认"))
    document.add_paragraph(_text(overall.interpretation, "待确认"))
    for block in overall.content_blocks or []:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(f"{_text(block.get('label'))}：{_text(block.get('text'))}")
        _apply_cjk_font(run)
        run.font.size = Pt(9)

    _add_heading(document, "四、监管要求清单")
    requirements = objects["requirements"]
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("条款", "规则类型", "责任主体", "行为/对象", "原文与证据")):
        _set_cell_text(cell, label, bold=True, color=ACCENT)
    for requirement in requirements:
        cells = table.add_row().cells
        _set_cell_text(cells[0], requirement.article_id)
        _set_cell_text(cells[1], requirement.rule_type)
        _set_cell_text(cells[2], requirement.subject)
        _set_cell_text(cells[3], f"{_text(requirement.action, '—')} / {_text(requirement.object, '—')}")
        evidence_ids = "、".join(evidence.evidence_id for evidence in requirement.evidence)
        _set_cell_text(cells[4], f"{_text(requirement.source_text)}\n证据：{evidence_ids or '待确认'}")

    _add_heading(document, "五、逐条解读")
    for interpretation in objects["article_interpretations"]:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(7)
        run = heading.add_run(_text(interpretation.article_id))
        _apply_cjk_font(run)
        run.bold = True
        run.font.color.rgb = ACCENT
        document.add_paragraph(_text(interpretation.summary))
        document.add_paragraph(_text(interpretation.interpretation))
        for block in interpretation.content_blocks or []:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(f"{_text(block.get('label'))}：{_text(block.get('text'))}")
            _apply_cjk_font(run)

    _add_heading(document, "六、版本比较")
    compare_table = document.add_table(rows=1, cols=2)
    compare_table.style = "Table Grid"
    compare_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, label in zip(compare_table.rows[0].cells, ("比较对象", "当前状态")):
        _set_cell_text(cell, label, bold=True, color=RGBColor(255, 255, 255), fill="181818")
    for label, value in (
        ("2015 年修订版（财金〔2015〕60号）", "待补充权威原文"),
        ("变化结论", "当前不生成新旧规差异结论"),
    ):
        cells = compare_table.add_row().cells
        _set_cell_text(cells[0], label, bold=True, color=ACCENT, fill="F7F7F8")
        _set_cell_text(cells[1], value)
    note = document.add_paragraph("只有在旧规原文、文件哈希和版本关系均完成核验后，才进入条款映射和变化解读。")
    note.paragraph_format.space_before = Pt(5)
    note.paragraph_format.space_after = Pt(8)
    _apply_cjk_font(note.runs[0])
    note.runs[0].font.color.rgb = RGBColor(128, 98, 39)

    _add_heading(document, "七、证据链与待确认边界")
    for evidence in objects["evidence"]:
        paragraph = document.add_paragraph(style="List Bullet")
        locator = evidence.locator or {}
        run = paragraph.add_run(
            f"{evidence.evidence_id}：{_text(evidence.source_text, '无原文片段')}；"
            f"位置：第{_text(locator.get('page'))}页，{_text(locator.get('article_no'))}；"
            f"核验状态：{_text(evidence.verification_status)}。"
        )
        _apply_cjk_font(run)
    document.add_paragraph("本报告仅基于已登记的监管原文生成。文号、附件、版本关系或适用性证据不足时，应回到工作台补充官方原文并重新复核；本报告不包含制度映射、差距分析、整改或审计结论。")

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {"file_name": output.name, "path": str(output), "content_version": 1}
