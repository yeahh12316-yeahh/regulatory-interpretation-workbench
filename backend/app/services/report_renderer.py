"""Render the locked Content Package into HTML and DOCX deliverables.

Both formats deliberately consume the same immutable payload.  The renderer
does not call a model or re-read mutable review objects, so a downloaded
report can always be traced back to one Content Package hash.
"""

from __future__ import annotations

import html
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(86, 125, 31)
INK = RGBColor(43, 49, 53)
# Hiragino is installed on macOS and is also discoverable by the DOCX renderer;
# using the family name avoids the missing-glyph boxes produced by STHeiti in
# headless LibreOffice checks.
CJK_FONT = "Hiragino Sans GB"


def _text(value: Any, fallback: str = "待确认") -> str:
    if value is None or value == "":
        return fallback
    if isinstance(value, list):
        return "；".join(str(item) for item in value) or fallback
    if isinstance(value, dict):
        return "；".join(f"{key}：{value[key]}" for key in value) or fallback
    return str(value)


def _e(value: Any, fallback: str = "待确认") -> str:
    return html.escape(_text(value, fallback))


def _package_rows(content: dict[str, Any]) -> list[str]:
    """Return stable, human-readable markers used by both renderers and QC."""
    regulation = content.get("regulation") or {}
    overview = content.get("overview") or {}
    rows = [
        _text(regulation.get("title")),
        _text(regulation.get("document_no")),
        _text(regulation.get("version_label")),
        _text(overview.get("summary")),
        _text(overview.get("interpretation")),
    ]
    for chapter in content.get("chapters") or []:
        interpretation = chapter.get("interpretation") or {}
        rows.extend([
            _text(chapter.get("article_no")),
            _text(chapter.get("original_text")),
            _text(interpretation.get("summary")),
            _text(interpretation.get("interpretation")),
        ])
        for requirement in chapter.get("requirements") or []:
            rows.extend([
                _text(requirement.get("requirement_id")),
                _text(requirement.get("source_text")),
                _text(requirement.get("subject")),
                _text(requirement.get("action")),
                _text(requirement.get("object")),
            ])
    for evidence in content.get("evidence_links") or []:
        rows.extend([_text(evidence.get("evidence_id")), _text(evidence.get("source_text"))])
    return [row for row in rows if row and row != "待确认"]


def render_report_html(content: dict[str, Any], output_path: str | Path) -> dict[str, Any]:
    regulation = content.get("regulation") or {}
    overview = content.get("overview") or {}
    package_id = _e(content.get("package_id"))
    content_hash = _e(content.get("content_hash") or "未记录")
    chapters = content.get("chapters") or []
    evidence_links = content.get("evidence_links") or []
    requirements = [item for chapter in chapters for item in (chapter.get("requirements") or [])]
    s5 = content.get("s5") or {}
    s5_output = s5.get("output") or {}
    comparison_status = s5_output.get("comparison_status") or "待确认"

    chapter_html = []
    for chapter in chapters:
        interpretation = chapter.get("interpretation") or {}
        requirement_html = "".join(
            f"<tr><td>{_e(item.get('requirement_id'))}</td><td>{_e(item.get('rule_type'))}</td>"
            f"<td>{_e(item.get('subject'))}</td><td>{_e(item.get('action'))} / {_e(item.get('object'))}</td>"
            f"<td>{_e(item.get('source_text'))}</td></tr>"
            for item in chapter.get("requirements") or []
        ) or '<tr><td colspan="5" class="muted">本条款暂无结构化要求</td></tr>'
        chapter_html.append(
            f"<article class=\"article-card\" id=\"article-{_e(chapter.get('article_id'))}\">"
            f"<div class=\"article-kicker\">{_e(chapter.get('article_no'))} · 第 {_e(chapter.get('source_page'))} 页</div>"
            f"<h3>{_e(chapter.get('article_no'))}</h3>"
            f"<div class=\"source-text\"><strong>法规原文</strong><p>{_e(chapter.get('original_text'))}</p></div>"
            f"<div class=\"interpretation\"><strong>条款解读</strong><p>{_e(interpretation.get('summary'))}</p>"
            f"<p>{_e(interpretation.get('interpretation'))}</p></div>"
            f"<table><thead><tr><th>要求编号</th><th>类型</th><th>责任主体</th><th>行为 / 对象</th><th>证据原文</th></tr></thead>"
            f"<tbody>{requirement_html}</tbody></table></article>"
        )
    evidence_html = "".join(
        f"<li><strong>{_e(item.get('evidence_id'))}</strong> · {_e(item.get('source_text'))} "
        f"<span class=\"muted\">位置：{_e((item.get('locator') or {}).get('page'))}页；状态：{_e(item.get('verification_status'))}</span></li>"
        for item in evidence_links
    ) or '<li class="muted">暂无证据记录</li>'

    document = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{_e(regulation.get('title'))} · 外规解读报告</title>
<style>
:root{{--ink:#2b3135;--muted:#68747b;--accent:#567d1f;--line:#dfe4e7;--soft:#f5f7f7;--warn:#806227}}
*{{box-sizing:border-box}} body{{margin:0;background:#eef1f2;color:var(--ink);font:14px/1.7 -apple-system,BlinkMacSystemFont,"PingFang SC","Microsoft YaHei",sans-serif}}
.report{{max-width:1120px;margin:0 auto;background:#fff;min-height:100vh;padding:48px 64px}}
.eyebrow{{color:var(--accent);font-size:12px;letter-spacing:.12em;font-weight:700}} h1{{margin:8px 0 2px;font-size:28px}} h2{{border-left:4px solid #86bc25;padding-left:10px;margin:36px 0 14px;font-size:20px}} h3{{margin:4px 0 12px;font-size:18px}}
.meta{{color:var(--muted);font-size:12px}} .package{{float:right;text-align:right;color:var(--muted);font-size:11px}} .meta-grid{{display:grid;grid-template-columns:150px 1fr;max-width:760px;border-top:1px solid var(--line)}} .meta-grid div{{padding:8px 10px;border-bottom:1px solid var(--line)}} .meta-grid .label{{background:var(--soft);font-weight:700;color:var(--accent)}}
.notice{{padding:14px 16px;background:#fff9ec;border:1px solid #ecd9aa;color:var(--warn);border-radius:6px}} .article-card{{border:1px solid var(--line);border-radius:8px;padding:20px;margin:16px 0;break-inside:avoid}} .article-kicker{{font-size:12px;color:var(--accent);font-weight:700}} .source-text{{background:#f8faf9;border-left:3px solid #b7c8ad;padding:10px 14px;margin:12px 0}} .interpretation{{padding:10px 14px;border-left:3px solid #86bc25;background:#fbfcf8}}
table{{width:100%;border-collapse:collapse;margin-top:14px;font-size:12px}} th,td{{border:1px solid var(--line);padding:7px 8px;text-align:left;vertical-align:top}} th{{background:#1b1d1f;color:#fff}} .muted{{color:var(--muted)}} li{{margin:8px 0}} footer{{border-top:1px solid var(--line);margin-top:40px;padding-top:12px;color:var(--muted);font-size:11px}}
@media(max-width:760px){{.report{{padding:28px 18px}} .package{{float:none;text-align:left;margin-top:12px}} .meta-grid{{grid-template-columns:110px 1fr}} table{{font-size:11px;display:block;overflow:auto}}}}
</style></head><body><main class="report" data-content-hash="{content_hash}" data-package-id="{package_id}">
<div class="package">Content Package<br>{package_id}<br>SHA-256 {content_hash}</div><div class="eyebrow">外规解读智能体工作台</div>
<h1>{_e(regulation.get('title'))}</h1><div class="meta">仅基于已登记监管原文 · 人工锁定版本</div>
<h2>一、法规概览与适用性</h2><div class="meta-grid">
<div class="label">文号</div><div>{_e(regulation.get('document_no'))}</div><div class="label">发布机关</div><div>{_e(regulation.get('issuer'))}</div>
<div class="label">版本</div><div>{_e(regulation.get('version_label'))}</div><div class="label">发布日期</div><div>{_e(regulation.get('publish_date'))}</div>
<div class="label">生效日期</div><div>{_e(regulation.get('effective_date'))}</div></div>
<div class="interpretation"><strong>整体解读</strong><p>{_e(overview.get('summary'))}</p><p>{_e(overview.get('interpretation'))}</p></div>
<h2>二、监管要求清单（{len(requirements)} 条）</h2><table><thead><tr><th>要求编号</th><th>类型</th><th>责任主体</th><th>行为 / 对象</th><th>证据原文</th></tr></thead><tbody>
{"".join(f"<tr><td>{_e(item.get('requirement_id'))}</td><td>{_e(item.get('rule_type'))}</td><td>{_e(item.get('subject'))}</td><td>{_e(item.get('action'))} / {_e(item.get('object'))}</td><td>{_e(item.get('source_text'))}</td></tr>" for item in requirements) or '<tr><td colspan="5" class="muted">暂无监管要求</td></tr>'}</tbody></table>
<h2>三、逐条解读</h2>{''.join(chapter_html) or '<p class="muted">暂无条款解读</p>'}
<h2>四、版本比较</h2><div class="notice">当前 S5 状态：{_e(comparison_status)}。只有旧规原文、文件哈希和版本关系完成核验后，才生成差异结论；本报告不伪造比较结果。</div>
<h2>五、Evidence 链路</h2><ul>{evidence_html}</ul>
<footer>本报告由锁定 Content Package 生成；上游内容变化后旧版本会标记为 STALE。报告不包含制度映射、差距分析、整改或审计结论。</footer>
</main></body></html>"""
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(document, encoding="utf-8")
    return {"file_name": output.name, "path": str(output), "content_hash": content.get("content_hash"), "package_id": content.get("package_id")}


def _font(run, *, size: int = 9, bold: bool = False, color: RGBColor = INK) -> None:
    run.font.name = CJK_FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), CJK_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _cell(cell, value: Any, *, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    if fill:
        props = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        props.append(shading)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(_text(value))
    _font(run, bold=bold, color=ACCENT if bold else INK)


def build_content_package_docx(content: dict[str, Any], output_path: str | Path) -> dict[str, Any]:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    normal = document.styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(9.5)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), CJK_FONT)
    regulation = content.get("regulation") or {}
    overview = content.get("overview") or {}
    title = document.add_paragraph()
    run = title.add_run(_text(regulation.get("title")))
    _font(run, size=20, bold=True, color=ACCENT)
    subtitle = document.add_paragraph("外规解读智能体工作台 · 人工锁定 Content Package")
    _font(subtitle.runs[0], size=10, color=RGBColor(100, 110, 116))
    package_meta = document.add_paragraph(f"Package：{_text(content.get('package_id'))}  |  SHA-256：{_text(content.get('content_hash'))}")
    _font(package_meta.runs[0], size=8, color=RGBColor(100, 110, 116))

    def heading(text: str) -> None:
        paragraph = document.add_heading(text, level=1)
        for item in paragraph.runs:
            _font(item, size=13, bold=True, color=ACCENT)

    heading("一、法规概览与适用性")
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in (("文号", regulation.get("document_no")), ("发布机关", regulation.get("issuer")), ("版本", regulation.get("version_label")), ("发布日期", regulation.get("publish_date")), ("生效日期", regulation.get("effective_date"))):
        cells = table.add_row().cells
        _cell(cells[0], label, bold=True, fill="F5F7F7")
        _cell(cells[1], value)
    document.add_paragraph(_text(overview.get("summary")))
    document.add_paragraph(_text(overview.get("interpretation")))

    heading("二、监管要求清单")
    requirements = [item for chapter in content.get("chapters") or [] for item in (chapter.get("requirements") or [])]
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("要求编号", "类型", "责任主体", "行为 / 对象", "证据原文")):
        _cell(cell, label, bold=True, fill="1B1D1F")
    for item in requirements:
        cells = table.add_row().cells
        for cell, value in zip(cells, (item.get("requirement_id"), item.get("rule_type"), item.get("subject"), f"{_text(item.get('action'))} / {_text(item.get('object'))}", item.get("source_text"))):
            _cell(cell, value)

    heading("三、逐条解读")
    for chapter in content.get("chapters") or []:
        interpretation = chapter.get("interpretation") or {}
        paragraph = document.add_paragraph()
        _font(paragraph.add_run(_text(chapter.get("article_no"))), size=11, bold=True, color=ACCENT)
        document.add_paragraph(f"法规原文：{_text(chapter.get('original_text'))}")
        document.add_paragraph(f"条款解读：{_text(interpretation.get('summary'))} {_text(interpretation.get('interpretation'))}")

    heading("四、版本比较")
    s5 = content.get("s5") or {}
    status = (s5.get("output") or {}).get("comparison_status") or "待确认"
    document.add_paragraph("六、版本比较")
    document.add_paragraph(f"当前 S5 状态：{_text(status)}。待补充权威原文、文件哈希和版本关系完成核验后，才生成差异结论；本报告不伪造比较结果。")

    heading("五、Evidence 链路")
    for evidence in content.get("evidence_links") or []:
        document.add_paragraph(f"{_text(evidence.get('evidence_id'))}：{_text(evidence.get('source_text'))}；位置：第 {_text((evidence.get('locator') or {}).get('page'))} 页；核验状态：{_text(evidence.get('verification_status'))}", style="List Bullet")
    document.add_paragraph("本报告仅基于已登记的监管原文生成；上游变化后旧 Content Package 会标记为 STALE，需要重新审核和导出。")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {"file_name": output.name, "path": str(output), "content_hash": content.get("content_hash"), "package_id": content.get("package_id")}


def check_render_consistency(content: dict[str, Any], html_path: str | Path, docx_path: str | Path) -> dict[str, Any]:
    html_text = Path(html_path).read_text(encoding="utf-8")
    document = Document(docx_path)
    docx_text = "\n".join([paragraph.text for paragraph in document.paragraphs] + [cell.text for table in document.tables for row in table.rows for cell in row.cells])
    markers = _package_rows(content)
    html_missing = [marker for marker in markers if marker not in html_text]
    docx_missing = [marker for marker in markers if marker not in docx_text]
    expected_hash = str(content.get("content_hash") or "")
    hash_ok = expected_hash and expected_hash in html_text and expected_hash in docx_text
    status = "passed" if not html_missing and not docx_missing and hash_ok else "blocked"
    return {
        "status": status,
        "package_id": content.get("package_id"),
        "content_hash": expected_hash,
        "checked_marker_count": len(markers),
        "html_missing": html_missing,
        "docx_missing": docx_missing,
        "hash_match": bool(hash_ok),
    }
