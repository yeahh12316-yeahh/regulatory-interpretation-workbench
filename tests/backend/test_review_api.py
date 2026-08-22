import asyncio
from io import BytesIO

import httpx
from docx import Document
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from sqlalchemy import create_engine
from sqlalchemy.orm import Session
from sqlalchemy.pool import StaticPool

from backend.app.core.config import get_settings
from backend.app.db import model_registry  # noqa: F401
from backend.app.db.base import Base
from backend.app.db.session import get_db
from backend.app.main import app


def build_review_fixture_pdf(include_document_no: bool = True) -> bytes:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    stream = BytesIO()
    document = canvas.Canvas(stream, pagesize=(420, 594))
    document.setFont("STSong-Light", 11)
    lines = [
        "财政部关于印发《测试复核办法（2026年版）》的通知",
        *( ["财金〔2026〕8号"] if include_document_no else [] ),
        "时间：2026-01-15",
        "测试复核办法（2026年版）",
        "第一条 为规范金融企业管理，制定本办法。",
        "第二条 本办法适用于中国境内依法设立的金融企业。金融企业应当建立复核机制。",
        "本办法自2026年3月1日起施行。",
    ]
    y = 550
    for line in lines:
        document.drawString(24, y, line)
        y -= 20
    document.save()
    return stream.getvalue()


def test_review_update_preserves_evidence_and_writes_audit(tmp_path, monkeypatch):
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)

    def db_override():
        with Session(engine) as session:
            yield session

    async def request(method: str, path: str, **kwargs):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            return await client.request(method, path, **kwargs)

    monkeypatch.setenv("PRIVATE_MODE", "true")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()
    app.dependency_overrides[get_db] = db_override
    try:
        task = asyncio.run(request("POST", "/api/tasks", json={"task_id": "REVIEW_TASK", "task_name": "复核测试任务"}))
        assert task.status_code == 201, task.text
        imported = asyncio.run(
            request(
                "POST",
                "/api/regulations/import",
                files={"file": ("测试复核办法（2026年版）.pdf", build_review_fixture_pdf(), "application/pdf")},
                data={"task_id": "REVIEW_TASK", "version_label": "2026年版"},
            )
        )
        assert imported.status_code == 201, imported.text
        run = asyncio.run(
            request(
                "POST",
                "/api/tasks/REVIEW_TASK/interpret",
                json={"institution_type": "商业银行", "business_scope": ["复核"], "region": "中国境内"},
            )
        )
        assert run.status_code == 200, run.text
        first_requirement = run.json()["requirements"][0]
        first_interpretation = run.json()["article_interpretations"][0]

        update_requirement = asyncio.run(
            request(
                "PATCH",
                f"/api/tasks/REVIEW_TASK/review/requirements/{first_requirement['requirement_id']}",
                json={"subject": "经人工确认的金融企业", "review_status": "reviewed"},
            )
        )
        assert update_requirement.status_code == 200, update_requirement.text
        updated_requirement = update_requirement.json()
        assert updated_requirement["subject"] == "经人工确认的金融企业"
        assert updated_requirement["source_text"] == first_requirement["source_text"]

        update_interpretation = asyncio.run(
            request(
                "PATCH",
                f"/api/tasks/REVIEW_TASK/review/interpretations/{first_interpretation['interpretation_id']}",
                json={"interpretation": "经人工复核的条款解读。", "review_status": "reviewed", "human_lock": True},
            )
        )
        assert update_interpretation.status_code == 200, update_interpretation.text
        assert update_interpretation.json()["human_lock"] is True

        llm = asyncio.run(request("POST", "/api/tasks/REVIEW_TASK/review/llm"))
        assert llm.status_code == 200, llm.text
        assert llm.json()["status"] == "not_configured"

        returned = asyncio.run(request("POST", "/api/tasks/REVIEW_TASK/review/decision", json={"decision": "return", "reason": "补充人工核验"}))
        assert returned.status_code == 200, returned.text
        assert returned.json()["task_status"] == "reviewing"

        review = asyncio.run(request("GET", "/api/tasks/REVIEW_TASK/review"))
        assert review.status_code == 200, review.text
        assert review.json()["audit_log_count"] >= 2
        assert any(item["requirement_id"] == first_requirement["requirement_id"] for item in review.json()["requirements"])
    finally:
        app.dependency_overrides.clear()
        get_settings.cache_clear()


def test_qc_blocks_export_until_all_results_are_reviewed(tmp_path, monkeypatch):
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)

    def db_override():
        with Session(engine) as session:
            yield session

    async def request(method: str, path: str, **kwargs):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            return await client.request(method, path, **kwargs)

    monkeypatch.setenv("PRIVATE_MODE", "true")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()
    app.dependency_overrides[get_db] = db_override
    try:
        task = asyncio.run(request("POST", "/api/tasks", json={"task_id": "QC_TASK", "task_name": "QC测试任务"}))
        assert task.status_code == 201, task.text
        imported = asyncio.run(
            request(
                "POST",
                "/api/regulations/import",
                files={"file": ("测试复核办法（2026年版）.pdf", build_review_fixture_pdf(include_document_no=False), "application/pdf")},
                data={"task_id": "QC_TASK", "version_label": "2026年版"},
            )
        )
        assert imported.status_code == 201, imported.text
        run = asyncio.run(
            request(
                "POST",
                "/api/tasks/QC_TASK/interpret",
                json={"institution_type": "商业银行", "business_scope": ["复核"], "region": "中国境内"},
            )
        )
        assert run.status_code == 200, run.text

        qc = asyncio.run(request("POST", "/api/tasks/QC_TASK/review/qc"))
        assert qc.status_code == 200, qc.text
        assert qc.json()["status"] == "blocked"
        assert qc.json()["blocker_count"] > 0
        assert qc.json()["task_status"] == "reviewing"
        assert "UNRESOLVED_METADATA" in {item["code"] for item in qc.json()["blockers"]}

        export = asyncio.run(request("POST", "/api/tasks/QC_TASK/export/docx"))
        assert export.status_code == 409, export.text
    finally:
        app.dependency_overrides.clear()
        get_settings.cache_clear()


def test_bulk_review_marks_all_reviewable_objects_and_writes_audit(tmp_path, monkeypatch):
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)

    def db_override():
        with Session(engine) as session:
            yield session

    async def request(method: str, path: str, **kwargs):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            return await client.request(method, path, **kwargs)

    monkeypatch.setenv("PRIVATE_MODE", "true")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()
    app.dependency_overrides[get_db] = db_override
    try:
        task = asyncio.run(request("POST", "/api/tasks", json={"task_id": "BULK_REVIEW_TASK", "task_name": "批量复核测试任务"}))
        assert task.status_code == 201, task.text
        imported = asyncio.run(request(
            "POST",
            "/api/regulations/import",
            files={"file": ("测试复核办法（2026年版）.pdf", build_review_fixture_pdf(), "application/pdf")},
            data={"task_id": "BULK_REVIEW_TASK", "version_label": "2026年版"},
        ))
        assert imported.status_code == 201, imported.text
        interpreted = asyncio.run(request(
            "POST",
            "/api/tasks/BULK_REVIEW_TASK/interpret",
            json={"institution_type": "商业银行", "business_scope": ["复核"], "region": "中国境内"},
        ))
        assert interpreted.status_code == 200, interpreted.text

        bulk = asyncio.run(request("POST", "/api/tasks/BULK_REVIEW_TASK/review/bulk"))

        assert bulk.status_code == 200, bulk.text
        payload = bulk.json()
        assert payload["review_summary"]["reviewed_requirements"] == payload["review_summary"]["total_requirements"]
        assert payload["review_summary"]["locked_interpretations"] == payload["review_summary"]["total_interpretations"]
        assert payload["review_summary"]["verified_evidence"] == payload["review_summary"]["total_evidence"]
        assert payload["audit_log_count"] >= payload["review_summary"]["total_requirements"] + payload["review_summary"]["total_interpretations"]
    finally:
        app.dependency_overrides.clear()
        get_settings.cache_clear()


def test_qc_pass_generates_a_downloadable_docx(tmp_path, monkeypatch):
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)

    def db_override():
        with Session(engine) as session:
            yield session

    async def request(method: str, path: str, **kwargs):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            return await client.request(method, path, **kwargs)

    monkeypatch.setenv("PRIVATE_MODE", "true")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()
    app.dependency_overrides[get_db] = db_override
    try:
        task = asyncio.run(request("POST", "/api/tasks", json={"task_id": "EXPORT_TASK", "task_name": "导出测试任务"}))
        assert task.status_code == 201, task.text
        imported = asyncio.run(
            request(
                "POST",
                "/api/regulations/import",
                files={"file": ("测试复核办法（2026年版）.pdf", build_review_fixture_pdf(), "application/pdf")},
                data={"task_id": "EXPORT_TASK", "version_label": "2026年版"},
            )
        )
        assert imported.status_code == 201, imported.text
        run = asyncio.run(
            request(
                "POST",
                "/api/tasks/EXPORT_TASK/interpret",
                json={"institution_type": "商业银行", "business_scope": ["复核"], "region": "中国境内"},
            )
        )
        assert run.status_code == 200, run.text
        payload = run.json()
        for requirement in payload["requirements"]:
            response = asyncio.run(
                request(
                    "PATCH",
                    f"/api/tasks/EXPORT_TASK/review/requirements/{requirement['requirement_id']}",
                    json={"review_status": "reviewed"},
                )
            )
            assert response.status_code == 200, response.text
        for interpretation in [payload["overall"], *payload["article_interpretations"]]:
            response = asyncio.run(
                request(
                    "PATCH",
                    f"/api/tasks/EXPORT_TASK/review/interpretations/{interpretation['interpretation_id']}",
                    json={"review_status": "reviewed", "human_lock": True},
                )
            )
            assert response.status_code == 200, response.text
        for evidence in payload["evidence"]:
            response = asyncio.run(
                request(
                    "PATCH",
                    f"/api/tasks/EXPORT_TASK/review/evidence/{evidence['evidence_id']}",
                    json={"verification_status": "verified"},
                )
            )
            assert response.status_code == 200, response.text

        qc = asyncio.run(request("POST", "/api/tasks/EXPORT_TASK/review/qc"))
        assert qc.status_code == 200, qc.text
        assert qc.json()["status"] == "passed", qc.text
        export = asyncio.run(request("POST", "/api/tasks/EXPORT_TASK/export/docx"))
        assert export.status_code == 200, export.text
        assert export.json()["review_status"] == "exported"
        assert export.json()["consistency"]["status"] == "passed"
        assert export.json()["html_download_url"].endswith("/html")
        download = asyncio.run(request("GET", export.json()["download_url"]))
        assert download.status_code == 200
        assert download.headers["content-type"].startswith("application/vnd.openxmlformats")
        assert download.content[:2] == b"PK"
        rendered = Document(BytesIO(download.content))
        rendered_text = "\n".join([paragraph.text for paragraph in rendered.paragraphs] + [cell.text for table in rendered.tables for row in table.rows for cell in row.cells])
        assert "六、版本比较" in rendered_text
        assert "待补充权威原文" in rendered_text
        html_download = asyncio.run(request("GET", export.json()["html_download_url"]))
        assert html_download.status_code == 200
        assert html_download.headers["content-type"].startswith("text/html")
        assert export.json()["consistency"]["content_hash"] in html_download.text

        package = asyncio.run(request("POST", "/api/tasks/EXPORT_TASK/content-package"))
        assert package.status_code == 201, package.text
        published = asyncio.run(request("POST", "/api/tasks/EXPORT_TASK/review/decision", json={"decision": "publish", "reason": "测试发布闸门"}))
        assert published.status_code == 200, published.text
        assert published.json()["task_status"] == "published"
    finally:
        app.dependency_overrides.clear()
        get_settings.cache_clear()


def test_s1_metadata_has_field_provenance_and_manual_override_survives_rerun(tmp_path, monkeypatch):
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)

    def db_override():
        with Session(engine) as session:
            yield session

    async def request(method: str, path: str, **kwargs):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            return await client.request(method, path, **kwargs)

    monkeypatch.setenv("PRIVATE_MODE", "true")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    get_settings.cache_clear()
    app.dependency_overrides[get_db] = db_override
    try:
        task = asyncio.run(request("POST", "/api/tasks", json={"task_id": "S1_METADATA_TASK", "task_name": "S1元数据闭环测试"}))
        assert task.status_code == 201, task.text
        imported = asyncio.run(
            request(
                "POST",
                "/api/regulations/import",
                files={"file": ("测试复核办法（2026年版）.pdf", build_review_fixture_pdf(include_document_no=False), "application/pdf")},
                data={"task_id": "S1_METADATA_TASK", "version_label": "2026年版"},
            )
        )
        assert imported.status_code == 201, imported.text
        run_payload = {
            "institution_type": "商业银行",
            "business_scope": ["复核"],
            "region": "中国境内",
        }
        run = asyncio.run(request("POST", "/api/tasks/S1_METADATA_TASK/interpret", json=run_payload))
        assert run.status_code == 200, run.text
        first_s1 = run.json()["stages"]["S1"]["output"]
        assert first_s1["metadata_fields"]["document_no"]["status"] == "missing"
        assert first_s1["metadata_fields"]["document_no"]["source_document_id"]
        assert first_s1["metadata_fields"]["document_no"]["extraction_method"] == "pypdf"

        updated = asyncio.run(
            request(
                "PATCH",
                "/api/tasks/S1_METADATA_TASK/review/metadata",
                json={
                    "document_no": "财金〔2026〕8号",
                    "issuer": ["财政部"],
                    "publish_date": "2026-01-15",
                    "effective_date": "2026-03-01",
                },
            )
        )
        assert updated.status_code == 200, updated.text
        updated_s1 = updated.json()["stages"]["S1"]["output"]
        assert updated_s1["metadata_fields"]["document_no"]["status"] == "manual_verified"
        assert updated_s1["metadata_fields"]["document_no"]["value"] == "财金〔2026〕8号"
        assert updated_s1["manual_overrides"]["document_no"]["value"] == "财金〔2026〕8号"
        assert "document_no" not in updated_s1["unresolved_fields"]

        rerun = asyncio.run(request("POST", "/api/tasks/S1_METADATA_TASK/interpret", json=run_payload))
        assert rerun.status_code == 200, rerun.text
        rerun_s1 = rerun.json()["stages"]["S1"]["output"]
        assert rerun_s1["document_no"] == "财金〔2026〕8号"
        assert rerun_s1["metadata_fields"]["document_no"]["status"] == "manual_verified"
        assert rerun_s1["metadata_fields"]["effective_date"]["status"] == "manual_verified"
    finally:
        app.dependency_overrides.clear()
        get_settings.cache_clear()
